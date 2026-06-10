"""Generate downloadable Word (.docx) and PDF documents from a structured outline.

Exposed to the model as the ``create_document`` function-calling tool: the model
passes a title, optional subtitle and an ordered list of ``sections`` (each with a
heading, body paragraphs, bullet points and an optional table). We render a real,
polished ``.docx`` (python-docx) or ``.pdf`` (reportlab) into ``.generated/{id}.{ext}``
and hand back a ``downloadUrl`` the model surfaces to the user as a Markdown link —
the same download plumbing as ``create_pptx``.

Both renderers share one content model and a deterministic accent theme keyed off the
title, so the same outline produces a consistent, modern-looking document in either
format. Heavy libraries are imported lazily inside the renderers so importing this
module never fails when they are absent.
"""

from __future__ import annotations

import hashlib
import re
from typing import Any

from deepseek_infra.core.errors import AppError, ErrorCode
from deepseek_infra.infra.tool_runtime.generated_files import store_generated_file

MAX_SECTIONS = 40
MAX_PARAGRAPHS_PER_SECTION = 40
MAX_BULLETS_PER_SECTION = 40
MAX_TABLE_ROWS = 60
MAX_TABLE_COLS = 8
MAX_PARAGRAPH_CHARS = 4000
SUPPORTED_FORMATS = ("docx", "pdf")
_FORMAT_ALIASES = {"word": "docx", "doc": "docx", "docx": "docx", "pdf": "pdf"}

# 一组精炼的专业配色：按标题哈希确定性选一套，让同一份内容稳定、不同主题之间有区分。
# primary 标题/强调、dark 大标题深色、light 表头底色、rule 分隔线浅色。
_DOC_THEMES = (
    {"primary": "1D4ED8", "dark": "1E3A8A", "light": "DBEAFE", "rule": "93C5FD"},  # 蓝
    {"primary": "4F46E5", "dark": "312E81", "light": "E0E7FF", "rule": "A5B4FC"},  # 靛
    {"primary": "0D9488", "dark": "115E59", "light": "CCFBF1", "rule": "5EEAD4"},  # 青
    {"primary": "EA580C", "dark": "9A3412", "light": "FFEDD5", "rule": "FDBA74"},  # 橙
    {"primary": "0284C7", "dark": "0C4A6E", "light": "E0F2FE", "rule": "7DD3FC"},  # 天蓝
    {"primary": "7C3AED", "dark": "581C87", "light": "F3E8FF", "rule": "C4B5FD"},  # 紫
)
_INK = "1F2933"        # 正文深灰
_MUTED = "6B7280"      # 次要文字 / 页脚
_BAND = "F1F5F9"       # 表格隔行底色
_CN_FONT = "微软雅黑"   # Word 中文字体
_PDF_FONT = "STSong-Light"  # reportlab 内置 Adobe 中文 CID 字体，无需附带字体文件

_pdf_font_ready = False


# ---------------------------------------------------------------------------
# 公共内容模型（与具体格式无关）
# ---------------------------------------------------------------------------


def create_document(fmt: str, title: str, sections: Any, *, subtitle: str = "") -> dict[str, Any]:
    """根据结构化大纲生成 Word 或 PDF，返回含 downloadUrl 的结果字典。"""
    normalized_fmt = _normalize_format(fmt)
    clean_title = str(title or "").strip()
    if not clean_title:
        raise AppError("文档需要一个标题（title）。", code=ErrorCode.INVALID_PAYLOAD)
    normalized_sections = _normalize_sections(sections)
    clean_subtitle = _clean_text(subtitle, limit=300)
    if normalized_fmt == "docx":
        return _create_word_document(clean_title, clean_subtitle, normalized_sections)
    return _create_pdf_document(clean_title, clean_subtitle, normalized_sections)


def _normalize_format(fmt: str) -> str:
    value = str(fmt or "").strip().lower()
    if value not in _FORMAT_ALIASES:
        raise AppError("format 必须是 docx（Word）或 pdf。", code=ErrorCode.INVALID_PAYLOAD)
    return _FORMAT_ALIASES[value]


def _clean_text(value: Any, *, limit: int = MAX_PARAGRAPH_CHARS) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()[:limit]


def _normalize_str_list(value: Any, limit: int) -> list[str]:
    if not isinstance(value, list):
        return []
    out: list[str] = []
    for item in value:
        text = _clean_text(item)
        if text:
            out.append(text)
        if len(out) >= limit:
            break
    return out


def _normalize_table(value: Any) -> dict[str, Any] | None:
    if not isinstance(value, dict):
        return None
    raw_headers = value.get("headers")
    raw_rows = value.get("rows")
    headers = [_clean_text(cell, limit=160) for cell in raw_headers][:MAX_TABLE_COLS] if isinstance(raw_headers, list) else []
    rows: list[list[str]] = []
    if isinstance(raw_rows, list):
        for raw_row in raw_rows[:MAX_TABLE_ROWS]:
            if not isinstance(raw_row, list):
                continue
            cells = [_clean_text(cell, limit=300) for cell in raw_row][:MAX_TABLE_COLS]
            if any(cells):
                rows.append(cells)
    width = max([len(headers), *(len(row) for row in rows)] or [0])
    if width == 0:
        return None
    if headers:
        headers = (headers + [""] * width)[:width]
    rows = [(row + [""] * width)[:width] for row in rows]
    return {"headers": headers, "rows": rows}


def _normalize_sections(sections: Any) -> list[dict[str, Any]]:
    if not isinstance(sections, list) or not sections:
        raise AppError("文档至少需要一个章节（sections）。", code=ErrorCode.INVALID_PAYLOAD)
    normalized: list[dict[str, Any]] = []
    for item in sections[:MAX_SECTIONS]:
        if not isinstance(item, dict):
            continue
        heading = _clean_text(item.get("heading"), limit=200)
        body = _normalize_str_list(item.get("body"), MAX_PARAGRAPHS_PER_SECTION)
        bullets = _normalize_str_list(item.get("bullets"), MAX_BULLETS_PER_SECTION)
        table = _normalize_table(item.get("table"))
        if not (heading or body or bullets or table):
            continue
        normalized.append({"heading": heading or "正文", "body": body, "bullets": bullets, "table": table})
    if not normalized:
        raise AppError("没有解析到有效的文档内容。", code=ErrorCode.INVALID_PAYLOAD)
    return normalized


def _theme_for(title: str) -> dict[str, str]:
    digest = hashlib.md5(str(title or "").encode("utf-8"), usedforsecurity=False).hexdigest()
    return _DOC_THEMES[int(digest, 16) % len(_DOC_THEMES)]


def _xml_escape(text: str) -> str:
    return str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _document_result(stored: dict[str, Any], *, fmt: str, title: str, sections: list[dict[str, Any]]) -> dict[str, Any]:
    outline = [
        {
            "index": index + 1,
            "heading": section["heading"],
            "paragraphs": len(section["body"]),
            "bullets": len(section["bullets"]),
            "hasTable": bool(section["table"]),
        }
        for index, section in enumerate(sections)
    ]
    label = "Word 文档" if fmt == "docx" else "PDF 文档"
    return {
        "fileId": stored["fileId"],
        "filename": stored["filename"],
        "format": fmt,
        "sectionCount": len(sections),
        "downloadUrl": stored["downloadUrl"],
        "title": title,
        "outline": outline,
        "note": (
            f"{label}已生成。请在最终回复里：先一句话说明文档标题、格式（{label}）和包含的主要章节，"
            "再用 Markdown 链接 [下载文档](downloadUrl) 把下载交给用户（链接 6 小时内有效）。"
            "不要把整篇正文重新粘回聊天，简述结构即可。"
        ),
    }


# ---------------------------------------------------------------------------
# Word (.docx) 渲染
# ---------------------------------------------------------------------------


def _create_word_document(title: str, subtitle: str, sections: list[dict[str, Any]]) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise AppError(
            "Word 生成需要 python-docx，请先安装：pip install python-docx",
            code=ErrorCode.INVALID_PAYLOAD,
        ) from exc

    theme = _theme_for(title)

    def style_run(run: Any, *, size: float, color: str, bold: bool = False) -> None:
        run.font.name = _CN_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _CN_FONT)

    def bottom_border(paragraph: Any, color: str, *, sz: int) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)
        p_pr.append(borders)

    def shade_cell(cell: Any, color: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shd)

    def add_page_field(paragraph: Any) -> Any:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)
        return run

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = _CN_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(_INK)
    try:
        normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _CN_FONT)
    except (AttributeError, KeyError):  # pragma: no cover - 取决于模板内部结构
        pass

    section_fmt = doc.sections[0]
    section_fmt.top_margin = Cm(2.2)
    section_fmt.bottom_margin = Cm(2.2)
    section_fmt.left_margin = Cm(2.4)
    section_fmt.right_margin = Cm(2.4)

    footer_paragraph = section_fmt.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer_paragraph.add_run("第 "), size=9, color=_MUTED)
    style_run(add_page_field(footer_paragraph), size=9, color=_MUTED)
    style_run(footer_paragraph.add_run(" 页"), size=9, color=_MUTED)

    # 标题块
    title_paragraph = doc.add_paragraph()
    style_run(title_paragraph.add_run(title), size=24, color=theme["dark"], bold=True)
    title_paragraph.paragraph_format.space_after = Pt(2)
    if subtitle:
        subtitle_paragraph = doc.add_paragraph()
        style_run(subtitle_paragraph.add_run(subtitle), size=12.5, color=_MUTED)
        subtitle_paragraph.paragraph_format.space_after = Pt(4)
    accent_rule = doc.add_paragraph()
    bottom_border(accent_rule, theme["primary"], sz=18)
    accent_rule.paragraph_format.space_after = Pt(10)

    for index, section in enumerate(sections, start=1):
        heading_paragraph = doc.add_paragraph(style="Heading 1")
        style_run(heading_paragraph.add_run(f"{index}. {section['heading']}"), size=15, color=theme["dark"], bold=True)
        heading_paragraph.paragraph_format.space_before = Pt(12)
        heading_paragraph.paragraph_format.space_after = Pt(4)
        bottom_border(heading_paragraph, theme["rule"], sz=8)

        for paragraph_text in section["body"]:
            body_paragraph = doc.add_paragraph()
            style_run(body_paragraph.add_run(paragraph_text), size=10.5, color=_INK)
            body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_paragraph.paragraph_format.space_after = Pt(6)
            body_paragraph.paragraph_format.line_spacing = 1.3

        for bullet_text in section["bullets"]:
            bullet_paragraph = doc.add_paragraph(style="List Bullet")
            style_run(bullet_paragraph.add_run(bullet_text), size=10.5, color=_INK)
            bullet_paragraph.paragraph_format.space_after = Pt(3)

        table = section["table"]
        if table:
            _add_word_table(
                doc,
                table,
                theme=theme,
                style_run=style_run,
                shade_cell=shade_cell,
                alignment=WD_TABLE_ALIGNMENT.CENTER,
            )
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    stored = store_generated_file(title, "docx", lambda path: doc.save(str(path)))
    return _document_result(stored, fmt="docx", title=title, sections=sections)


def _add_word_table(doc: Any, table: dict[str, Any], *, theme: dict[str, str], style_run: Any, shade_cell: Any, alignment: Any) -> None:
    headers = table["headers"]
    rows = table["rows"]
    has_header = bool(headers)
    col_count = len(headers) if has_header else len(rows[0])
    row_count = len(rows) + (1 if has_header else 0)

    word_table = doc.add_table(rows=row_count, cols=col_count)
    try:
        word_table.style = "Table Grid"
    except KeyError:  # pragma: no cover - 模板缺少内置表样式时退化为无边框
        pass
    word_table.alignment = alignment

    offset = 0
    if has_header:
        for col, text in enumerate(headers):
            cell = word_table.cell(0, col)
            style_run(cell.paragraphs[0].add_run(text), size=10, color="FFFFFF", bold=True)
            shade_cell(cell, theme["primary"])
        offset = 1

    for row_index, row in enumerate(rows):
        for col, text in enumerate(row):
            cell = word_table.cell(offset + row_index, col)
            style_run(cell.paragraphs[0].add_run(text), size=10, color=_INK)
            if row_index % 2 == 1:
                shade_cell(cell, _BAND)


# ---------------------------------------------------------------------------
# PDF 渲染
# ---------------------------------------------------------------------------


def _create_pdf_document(title: str, subtitle: str, sections: list[dict[str, Any]]) -> dict[str, Any]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            HRFlowable,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ModuleNotFoundError as exc:
        raise AppError(
            "PDF 生成需要 reportlab，请先安装：pip install reportlab",
            code=ErrorCode.INVALID_PAYLOAD,
        ) from exc

    _register_cjk_font(pdfmetrics, UnicodeCIDFont)
    theme = _theme_for(title)
    primary = colors.HexColor(f"#{theme['primary']}")
    dark = colors.HexColor(f"#{theme['dark']}")
    rule = colors.HexColor(f"#{theme['rule']}")
    ink = colors.HexColor(f"#{_INK}")
    muted = colors.HexColor(f"#{_MUTED}")
    band = colors.HexColor(f"#{_BAND}")

    title_style = ParagraphStyle("DocTitle", fontName=_PDF_FONT, fontSize=24, leading=29, textColor=dark, spaceAfter=2)
    subtitle_style = ParagraphStyle("DocSubtitle", fontName=_PDF_FONT, fontSize=12.5, leading=17, textColor=muted, spaceAfter=2)
    heading_style = ParagraphStyle(
        "DocHeading", fontName=_PDF_FONT, fontSize=15, leading=20, textColor=dark, spaceBefore=12, spaceAfter=2, keepWithNext=1
    )
    body_style = ParagraphStyle("DocBody", fontName=_PDF_FONT, fontSize=10.5, leading=16.5, textColor=ink, alignment=TA_JUSTIFY, spaceAfter=6)
    cell_style = ParagraphStyle("DocCell", fontName=_PDF_FONT, fontSize=9.5, leading=13, textColor=ink)
    head_cell_style = ParagraphStyle("DocHeadCell", fontName=_PDF_FONT, fontSize=9.5, leading=13, textColor=colors.white)

    story: list[Any] = [Paragraph(_xml_escape(title), title_style)]
    if subtitle:
        story.append(Paragraph(_xml_escape(subtitle), subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=primary, spaceBefore=4, spaceAfter=12))

    page_width = A4[0] - 40 * mm  # 左右各 20mm 边距
    for index, section in enumerate(sections, start=1):
        story.append(Paragraph(f"{index}. {_xml_escape(section['heading'])}", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.6, color=rule, spaceBefore=2, spaceAfter=8))
        for paragraph_text in section["body"]:
            story.append(Paragraph(_xml_escape(paragraph_text), body_style))
        if section["bullets"]:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_xml_escape(bullet), body_style), spaceAfter=2) for bullet in section["bullets"]],
                    bulletType="bullet",
                    start="•",
                    bulletColor=primary,
                    # STSong CID 字体没有 • 字形，bullet 符号改用始终可用的 Helvetica 渲染。
                    bulletFontName="Helvetica",
                    bulletFontSize=9,
                    leftIndent=16,
                )
            )
        table = section["table"]
        if table:
            story.append(Spacer(1, 4))
            story.append(
                _build_pdf_table(
                    table,
                    page_width=page_width,
                    cell_style=cell_style,
                    head_cell_style=head_cell_style,
                    header_bg=primary,
                    grid=rule,
                    band=band,
                    paragraph_cls=Paragraph,
                    table_cls=Table,
                    table_style_cls=TableStyle,
                )
            )
        story.append(Spacer(1, 12))

    def decorate(canvas: Any, doc_template: Any, *, later: bool) -> None:
        canvas.saveState()
        width, height = A4
        canvas.setStrokeColor(rule)
        canvas.setLineWidth(0.5)
        canvas.line(doc_template.leftMargin, 15 * mm, width - doc_template.rightMargin, 15 * mm)
        canvas.setFont(_PDF_FONT, 8.5)
        canvas.setFillColor(muted)
        canvas.drawCentredString(width / 2, 10.5 * mm, f"第 {canvas.getPageNumber()} 页")
        if later:
            canvas.drawString(doc_template.leftMargin, height - 14 * mm, title[:60])
            canvas.setStrokeColor(rule)
            canvas.setLineWidth(0.4)
            canvas.line(doc_template.leftMargin, height - 15.5 * mm, width - doc_template.rightMargin, height - 15.5 * mm)
        canvas.restoreState()

    def writer(path: Any) -> None:
        pdf = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=22 * mm,
            bottomMargin=20 * mm,
            title=title,
        )
        pdf.build(
            list(story),
            onFirstPage=lambda canvas, d: decorate(canvas, d, later=False),
            onLaterPages=lambda canvas, d: decorate(canvas, d, later=True),
        )

    stored = store_generated_file(title, "pdf", writer)
    return _document_result(stored, fmt="pdf", title=title, sections=sections)


def _build_pdf_table(
    table: dict[str, Any],
    *,
    page_width: float,
    cell_style: Any,
    head_cell_style: Any,
    header_bg: Any,
    grid: Any,
    band: Any,
    paragraph_cls: Any,
    table_cls: Any,
    table_style_cls: Any,
) -> Any:
    headers = table["headers"]
    rows = table["rows"]
    has_header = bool(headers)
    col_count = len(headers) if has_header else len(rows[0])

    data: list[list[Any]] = []
    if has_header:
        data.append([paragraph_cls(_xml_escape(cell), head_cell_style) for cell in headers])
    for row in rows:
        data.append([paragraph_cls(_xml_escape(cell), cell_style) for cell in row])

    col_width = page_width / col_count
    pdf_table = table_cls(data, colWidths=[col_width] * col_count, repeatRows=1 if has_header else 0)

    commands: list[tuple[Any, ...]] = [
        ("GRID", (0, 0), (-1, -1), 0.5, grid),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    start = 0
    if has_header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), header_bg))
        start = 1
    for offset, row_index in enumerate(range(start, len(data))):
        if offset % 2 == 1:
            commands.append(("BACKGROUND", (0, row_index), (-1, row_index), band))
    pdf_table.setStyle(table_style_cls(commands))
    return pdf_table


def _register_cjk_font(pdfmetrics: Any, unicode_cid_font: Any) -> None:
    global _pdf_font_ready
    if _pdf_font_ready:
        return
    pdfmetrics.registerFont(unicode_cid_font(_PDF_FONT))
    _pdf_font_ready = True
